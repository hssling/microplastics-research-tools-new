from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
import matplotlib.pyplot as plt

# Create document
doc = Document()

# Set to landscape orientation
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
new_width, new_height = section.page_height, section.page_width
section.page_width = new_width
section.page_height = new_height

# Title
title = doc.add_heading("Dietary Guidelines for Indians - Infographic", 0)
title.alignment = 1

doc.add_paragraph(
    "A visual guide by the National Institute of Nutrition (NIN) "
    "to promote health, prevent disease, and ensure a well-nourished nation."
)

# --- Section 1: India's Dual Nutrition Challenge ---
doc.add_heading("India's Dual Nutrition Challenge", level=1)

doc.add_heading("Undernutrition", level=2)
data1 = {"Low Birth Weight": 22, "Underweight <5y": 43, "Stunting <5y": 48,
         "Wasting <5y": 20, "Anaemia (Women)": 75}
plt.bar(data1.keys(), data1.values(), color="#C62828")
plt.xticks(rotation=30, ha="right")
plt.ylabel("Prevalence (%)")
plt.title("Undernutrition in India")
plt.tight_layout()
plt.savefig("undernutrition.png")
plt.close()
doc.add_picture("undernutrition.png", width=Inches(4.5))

doc.add_heading("Overnutrition & NCDs", level=2)
data2 = {"Overweight Rural": 9.4, "Overweight Urban": 38, "Hypertension Rural": 25,
         "Hypertension Urban": 35, "Diabetes Urban": 16}
plt.bar(data2.keys(), data2.values(), color="#283593")
plt.xticks(rotation=30, ha="right")
plt.ylabel("Prevalence (%)")
plt.title("Overnutrition in India")
plt.tight_layout()
plt.savefig("overnutrition.png")
plt.close()
doc.add_picture("overnutrition.png", width=Inches(4.5))

# --- Section 2: Balanced Diet ---
doc.add_heading("Guideline 1: The Foundation of a Balanced Diet", level=1)
doc.add_paragraph("A balanced diet provides all necessary nutrients in required amounts "
                  "and proper proportions. Achieved by eating a variety of foods.")

# Macronutrient chart
labels = ["Carbohydrates", "Proteins", "Fats"]
sizes = [55, 15, 30]
colors = ["#00838F", "#FFAB00", "#AD1457"]
plt.pie(sizes, labels=labels, autopct="%1.0f%%", colors=colors)
plt.title("Macronutrient Distribution")
plt.savefig("macronutrients.png")
plt.close()
doc.add_picture("macronutrients.png", width=Inches(4.5))

# Food groups
doc.add_paragraph("🌾 Cereals, Millets & Pulses – Main source of energy and protein")
doc.add_paragraph("🥦 Vegetables & Fruits – Rich in vitamins, minerals, and fiber")
doc.add_paragraph("🥛 Milk, Eggs & Meat – High-quality protein and micronutrients")
doc.add_paragraph("🥑 Fats, Nuts & Oils – Energy and essential fatty acids")

# --- Section 3: Nutrition Through Life Stages ---
doc.add_heading("Nutrition Through Life Stages", level=1)

# Pregnancy table
table = doc.add_table(rows=3, cols=2)
table.style = "Light List"
table.cell(0,0).text, table.cell(0,1).text = "Life Stage", "Extra Kcal/day"
table.cell(1,0).text, table.cell(1,1).text = "Pregnancy", "+350"
table.cell(2,0).text, table.cell(2,1).text = "Lactation", "+600"

doc.add_paragraph("👶 Infancy (0–6m): Exclusive breastfeeding")
doc.add_paragraph("🍲 After 6m: Semi-solid complementary foods")
doc.add_paragraph("🧒 Childhood: Balanced meals, prevent deficiencies")
doc.add_paragraph("👩 Adolescents: Extra calcium & iron")
doc.add_paragraph("👵 Elderly: Nutrient-dense, soft, easy-to-digest meals")

# --- Section 4: Healthy Habits ---
doc.add_heading("Building Healthy Habits for Life", level=1)
doc.add_paragraph("🥦 Eat Vegetables & Fruits – Aim ≥ 300g/day vegetables & ≥ 100g/day fruits")
doc.add_paragraph("🧂 Limit Salt – <5g/day; avoid pickles, papads, chips")
doc.add_paragraph("🍳 Cook Smart – Prefer steaming/boiling, avoid reused oil")
doc.add_paragraph("💧 Hydrate – At least 8 glasses safe water daily")
doc.add_paragraph("🏃 Be Active – 45 min/day, 5+ days per week")

# Footer
doc.add_paragraph("\nInfographic based on the 'Dietary Guidelines for Indians - A Manual' "
                  "by NIN-ICMR, Hyderabad.\nThis is a visual interpretation and not a "
                  "replacement for the official document.")


# Enhanced error handling and logging for file save
import os
try:
    output_path = os.path.abspath("Dietary_Guidelines_Infographic.docx")
    doc.save(output_path)
    print(f"Word infographic created: {output_path}")
except Exception as e:
    print(f"Error saving file: {e}")

