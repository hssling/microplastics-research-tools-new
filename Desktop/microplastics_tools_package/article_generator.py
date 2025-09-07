import openai
from typing import Dict, Any, List
import json
from datetime import datetime

class ArticleGenerator:
    def __init__(self, api_key: str):
        self.client = openai.OpenAI(api_key=api_key)
    def __init__(self, api_key: str = None, ai_manager=None):
        self.ai_manager = ai_manager

    def generate_systematic_review_article(self, protocol: Dict[str, Any],
                                         studies: List[Dict[str, Any]],
                                         extracted_data: List[Dict[str, Any]],
                                         quality_assessments: List[Dict[str, Any]],
                                         meta_results: Dict[str, Any]) -> Dict[str, Any]:
                                         meta_results: Dict[str, Any],
                                         model: str) -> Dict[str, Any]:
        """
        Generate a complete systematic review article
        """
        article_sections = {}

        # Generate each section
        article_sections['title'] = self._generate_title(protocol)
        article_sections['abstract'] = self._generate_abstract(protocol, studies, meta_results)
        article_sections['introduction'] = self._generate_introduction(protocol)
        article_sections['methods'] = self._generate_methods(protocol, studies)
        article_sections['results'] = self._generate_results(studies, extracted_data, meta_results)
        article_sections['discussion'] = self._generate_discussion(meta_results, quality_assessments)
        article_sections['conclusion'] = self._generate_conclusion(meta_results)
        article_sections['title'] = self._generate_title(protocol, model)
        article_sections['abstract'] = self._generate_abstract(protocol, studies, meta_results, model)
        article_sections['introduction'] = self._generate_introduction(protocol, model)
        article_sections['methods'] = self._generate_methods(protocol, studies, model)
        article_sections['results'] = self._generate_results(studies, extracted_data, meta_results, model)
        article_sections['discussion'] = self._generate_discussion(meta_results, quality_assessments, model)
        article_sections['conclusion'] = self._generate_conclusion(meta_results, model)
        article_sections['references'] = self._generate_references(studies)

        # Add metadata
        article = {
            'metadata': {
                'generated_at': datetime.utcnow().isoformat(),
                'protocol_id': protocol.get('id'),
                'num_studies': len(studies),
                'ai_model': 'gpt-4'
                'ai_model': model
            },
            'sections': article_sections
        }

        return article

    def _generate_title(self, protocol: Dict[str, Any]) -> str:
    def _generate_title(self, protocol: Dict[str, Any], model: str) -> str:
        """
        Generate article title
        """
        research_question = protocol.get('research_question', '')
        
        prompt = f"""
        Generate a clear, concise title for a systematic review article based on this research question:
        {research_question}
        
        The title should follow the format: "Systematic Review: [Topic] - A Review of [Key Elements]"
        Make it academic and informative.
        """

        try:
            response = self.client.chat.completions.create(
                model="gpt-4",
            if not self.ai_manager:
                raise ValueError("AI Manager not configured.")

            title = self.ai_manager.call_model(
                messages=[
                    {"role": "system", "content": "You are an expert in academic writing. Generate clear, concise titles for systematic review articles."},
                    {"role": "user", "content": prompt}
                ],
                model=model,
                max_tokens=100,
                temperature=0.7
            )
            return title.strip().replace('"', '')

            return response.choices[0].message.content.strip()

        except Exception as e:
            return f"Systematic Review: {research_question[:50]}..."

    def _generate_abstract(self, protocol: Dict[str, Any], studies: List[Dict[str, Any]], 
                          meta_results: Dict[str, Any]) -> str:
                          meta_results: Dict[str, Any], model: str) -> str:
        """
        Generate structured abstract
        """
        research_question = protocol.get('research_question', '')
        num_studies = len(studies)
        
        # Extract key findings from meta results
        key_findings = ""
        if 'pooled_effect' in meta_results:
            effect = meta_results['pooled_effect']
            ci_lower = meta_results.get('confidence_interval_lower', 0)
            ci_upper = meta_results.get('confidence_interval_upper', 0)
            key_findings = f"The pooled effect size was {effect:.3f} (95% CI: {ci_lower:.3f} to {ci_upper:.3f})."

        prompt = f"""
        Generate a structured abstract for a systematic review with the following details:

        Research Question: {research_question}
        Number of Studies: {num_studies}
        Key Findings: {key_findings}

        Structure the abstract with these sections:
        - Background
        - Methods
        - Results
        - Conclusion

        Keep it concise (200-250 words).
        """

        try:
            response = self.client.chat.completions.create(
                model="gpt-4",
            if not self.ai_manager:
                raise ValueError("AI Manager not configured.")

            abstract = self.ai_manager.call_model(
                messages=[
                    {"role": "system", "content": "You are an expert in academic writing. Generate structured abstracts for systematic reviews following PRISMA guidelines."},
                    {"role": "user", "content": prompt}
                ],
                model=model,
                max_tokens=400,
                temperature=0.6
            )
            return abstract.strip()

            # Check if response has the expected structure
            if response and response.choices and len(response.choices) > 0:
                if hasattr(response.choices[0], 'message') and response.choices[0].message:
                    if hasattr(response.choices[0].message, 'content') and response.choices[0].message.content:
                        return response.choices[0].message.content.strip()

            # If we get here, the response structure is unexpected
            print(f"Unexpected API response structure in _generate_abstract: {response}")
            return f"Abstract generation failed: Unexpected API response structure"

        except Exception as e:
            print(f"OpenAI API error in _generate_abstract: {str(e)}")
            print(f"AI API error in _generate_abstract: {str(e)}")
            return f"Abstract generation failed: {str(e)}"

    def _generate_introduction(self, protocol: Dict[str, Any]) -> str:
    def _generate_introduction(self, protocol: Dict[str, Any], model: str) -> str:
        """
        Generate introduction section
        """
        research_question = protocol.get('research_question', '')
        background = protocol.get('background', '')
        objectives = protocol.get('objectives', '')

        prompt = f"""
        Generate the introduction section for a systematic review article.

        Research Question: {research_question}
        Background: {background}
        Objectives: {objectives}

        The introduction should include:
        1. Background and rationale
        2. Research question and objectives
        3. Brief overview of existing literature
        4. Importance of the review

        Write in academic style, approximately 400-600 words.
        """

        try:
            response = self.client.chat.completions.create(
                model="gpt-4",
            if not self.ai_manager:
                raise ValueError("AI Manager not configured.")

            introduction = self.ai_manager.call_model(
                messages=[
                    {"role": "system", "content": "You are an expert in academic writing. Generate comprehensive introduction sections for systematic reviews."},
                    {"role": "user", "content": prompt}
                ],
                model=model,
                max_tokens=800,
                temperature=0.7
            )
            return introduction.strip()

            # Check if response has the expected structure
            if response and response.choices and len(response.choices) > 0:
                if hasattr(response.choices[0], 'message') and response.choices[0].message:
                    if hasattr(response.choices[0].message, 'content') and response.choices[0].message.content:
                        return response.choices[0].message.content.strip()

            # If we get here, the response structure is unexpected
            print(f"Unexpected API response structure in _generate_introduction: {response}")
            return f"Introduction generation failed: Unexpected API response structure"

        except Exception as e:
            print(f"OpenAI API error in _generate_introduction: {str(e)}")
            print(f"AI API error in _generate_introduction: {str(e)}")
            return f"Introduction generation failed: {str(e)}"

    def _generate_methods(self, protocol: Dict[str, Any], studies: List[Dict[str, Any]]) -> str:
    def _generate_methods(self, protocol: Dict[str, Any], studies: List[Dict[str, Any]], model: str) -> str:
        """
        Generate methods section
        """
        inclusion_criteria = protocol.get('inclusion_criteria', '')
        exclusion_criteria = protocol.get('exclusion_criteria', '')
        search_strategy = protocol.get('search_strategy', '')
        num_studies = len(studies)

        prompt = f"""
        Generate the methods section for a systematic review article.

        Eligibility Criteria:
        Inclusion: {inclusion_criteria}
        Exclusion: {exclusion_criteria}

        Search Strategy: {search_strategy}
        Number of Studies Included: {num_studies}

        The methods section should include:
        1. Study design and registration
        2. Eligibility criteria
        3. Information sources and search strategy
        4. Study selection process
        5. Data extraction methods
        6. Risk of bias assessment
        7. Data synthesis methods

        Follow PRISMA guidelines for reporting.
        """

        try:
            response = self.client.chat.completions.create(
                model="gpt-4",
            if not self.ai_manager:
                raise ValueError("AI Manager not configured.")

            methods = self.ai_manager.call_model(
                messages=[
                    {"role": "system", "content": "You are an expert in systematic review methodology. Generate detailed methods sections following PRISMA guidelines."},
                    {"role": "user", "content": prompt}
                ],
                model=model,
                max_tokens=1000,
                temperature=0.6
            )
            return methods.strip()

            # Check if response has the expected structure
            if response and response.choices and len(response.choices) > 0:
                if hasattr(response.choices[0], 'message') and response.choices[0].message:
                    if hasattr(response.choices[0].message, 'content') and response.choices[0].message.content:
                        return response.choices[0].message.content.strip()

            # If we get here, the response structure is unexpected
            print(f"Unexpected API response structure in _generate_methods: {response}")
            return f"Methods generation failed: Unexpected API response structure"

        except Exception as e:
            print(f"OpenAI API error in _generate_methods: {str(e)}")
            print(f"AI API error in _generate_methods: {str(e)}")
            return f"Methods generation failed: {str(e)}"

    def _generate_results(self, studies: List[Dict[str, Any]], 
                         extracted_data: List[Dict[str, Any]], 
                         meta_results: Dict[str, Any]) -> str:
                         meta_results: Dict[str, Any], model: str) -> str:
        """
        Generate results section
        """
        num_studies = len(studies)
        
        # Summarize study characteristics
        study_summary = f"Total of {num_studies} studies were included in the review."
        
        # Extract key meta-analysis results
        meta_summary = ""
        if 'pooled_effect' in meta_results:
            effect = meta_results['pooled_effect']
            ci_lower = meta_results.get('confidence_interval_lower', 0)
            ci_upper = meta_results.get('confidence_interval_upper', 0)
            i2 = meta_results.get('heterogeneity_i2', 0)
            p_value = meta_results.get('p_value', 1.0)
            
            meta_summary = f"""
            The pooled effect size was {effect:.3f} (95% CI: {ci_lower:.3f} to {ci_upper:.3f}, 
            p = {p_value:.3f}). Heterogeneity was {i2:.1f}%.
            """

        prompt = f"""
        Generate the results section for a systematic review article.

        Study Summary: {study_summary}
        Meta-Analysis Results: {meta_summary}
        Number of Studies: {num_studies}

        The results section should include:
        1. Study selection flow (PRISMA diagram summary)
        2. Study characteristics
        3. Risk of bias assessment results
        4. Primary outcome results
        5. Secondary outcomes (if available)
        6. Meta-analysis results with forest plot description
        7. Heterogeneity assessment
        8. Subgroup analyses (if performed)

        Present results clearly with appropriate statistics.
        """

        try:
            response = self.client.chat.completions.create(
                model="gpt-4",
            if not self.ai_manager:
                raise ValueError("AI Manager not configured.")

            results = self.ai_manager.call_model(
                messages=[
                    {"role": "system", "content": "You are an expert in academic writing. Generate clear, comprehensive results sections for systematic reviews."},
                    {"role": "user", "content": prompt}
                ],
                model=model,
                max_tokens=1200,
                temperature=0.6
            )
            return results.strip()

            # Check if response has the expected structure
            if response and response.choices and len(response.choices) > 0:
                if hasattr(response.choices[0], 'message') and response.choices[0].message:
                    if hasattr(response.choices[0].message, 'content') and response.choices[0].message.content:
                        return response.choices[0].message.content.strip()

            # If we get here, the response structure is unexpected
            print(f"Unexpected API response structure in _generate_results: {response}")
            return f"Results generation failed: Unexpected API response structure"

        except Exception as e:
            print(f"OpenAI API error in _generate_results: {str(e)}")
            print(f"AI API error in _generate_results: {str(e)}")
            return f"Results generation failed: {str(e)}"

    def _generate_discussion(self, meta_results: Dict[str, Any], 
                           quality_assessments: List[Dict[str, Any]]) -> str:
                           quality_assessments: List[Dict[str, Any]], model: str) -> str:
        """
        Generate discussion section
        """
        # Summarize quality assessments
        quality_summary = self._summarize_quality_assessments(quality_assessments)
        
        # Extract key findings
        key_findings = ""
        if 'pooled_effect' in meta_results:
            effect = meta_results['pooled_effect']
            ci_lower = meta_results.get('confidence_interval_lower', 0)
            ci_upper = meta_results.get('confidence_interval_upper', 0)
            key_findings = f"The meta-analysis showed a pooled effect of {effect:.3f} (95% CI: {ci_lower:.3f} to {ci_upper:.3f})."

        prompt = f"""
        Generate the discussion section for a systematic review article.

        Key Findings: {key_findings}
        Quality Assessment Summary: {quality_summary}

        The discussion section should include:
        1. Summary of main findings
        2. Comparison with existing literature
        3. Strengths and limitations of the review
        4. Quality assessment implications
        5. Implications for practice/policy/research
        6. Future research directions

        Be balanced and evidence-based in interpretations.
        """

        try:
            response = self.client.chat.completions.create(
                model="gpt-4",
            if not self.ai_manager:
                raise ValueError("AI Manager not configured.")

            discussion = self.ai_manager.call_model(
                messages=[
                    {"role": "system", "content": "You are an expert in academic writing. Generate balanced, evidence-based discussion sections for systematic reviews."},
                    {"role": "user", "content": prompt}
                ],
                model=model,
                max_tokens=1000,
                temperature=0.7
            )
            return discussion.strip()

            # Check if response has the expected structure
            if response and response.choices and len(response.choices) > 0:
                if hasattr(response.choices[0], 'message') and response.choices[0].message:
                    if hasattr(response.choices[0].message, 'content') and response.choices[0].message.content:
                        return response.choices[0].message.content.strip()

            # If we get here, the response structure is unexpected
            print(f"Unexpected API response structure in _generate_discussion: {response}")
            return f"Discussion generation failed: Unexpected API response structure"

        except Exception as e:
            print(f"OpenAI API error in _generate_discussion: {str(e)}")
            print(f"AI API error in _generate_discussion: {str(e)}")
            return f"Discussion generation failed: {str(e)}"

    def _generate_conclusion(self, meta_results: Dict[str, Any]) -> str:
    def _generate_conclusion(self, meta_results: Dict[str, Any], model: str) -> str:
        """
        Generate conclusion section
        """
        key_findings = ""
        if 'pooled_effect' in meta_results:
            effect = meta_results['pooled_effect']
            key_findings = f"The review concludes that the intervention has a pooled effect size of {effect:.3f}."

        prompt = f"""
        Generate the conclusion section for a systematic review article.

        Key Findings: {key_findings}

        The conclusion should:
        1. Summarize the main findings
        2. Highlight clinical/practical implications
        3. Note any recommendations
        4. Suggest future research needs

        Keep it concise but comprehensive.
        """

        try:
            response = self.client.chat.completions.create(
                model="gpt-4",
            if not self.ai_manager:
                raise ValueError("AI Manager not configured.")

            conclusion = self.ai_manager.call_model(
                messages=[
                    {"role": "system", "content": "You are an expert in academic writing. Generate concise, impactful conclusions for systematic reviews."},
                    {"role": "user", "content": prompt}
                ],
                model=model,
                max_tokens=300,
                temperature=0.6
            )
            return conclusion.strip()

            # Check if response has the expected structure
            if response and response.choices and len(response.choices) > 0:
                if hasattr(response.choices[0], 'message') and response.choices[0].message:
                    if hasattr(response.choices[0].message, 'content') and response.choices[0].message.content:
                        return response.choices[0].message.content.strip()

            # If we get here, the response structure is unexpected
            print(f"Unexpected API response structure in _generate_conclusion: {response}")
            return f"Conclusion generation failed: Unexpected API response structure"

        except Exception as e:
            print(f"OpenAI API error in _generate_conclusion: {str(e)}")
            print(f"AI API error in _generate_conclusion: {str(e)}")
            return f"Conclusion generation failed: {str(e)}"

    def _generate_references(self, studies: List[Dict[str, Any]]) -> str:
        """
        Generate references section
        """
        references = []
        
        for study in studies:
            title = study.get('title', 'Unknown Title')
            authors = study.get('authors', [])
            journal = study.get('journal', 'Unknown Journal')
            year = study.get('year', 'Unknown Year')
            doi = study.get('doi', '')
            
            # Format as basic reference
            if authors:
                author_str = authors[0] if isinstance(authors, list) else str(authors)
                if len(authors) > 1:
                    author_str += " et al."
            else:
                author_str = "Unknown Authors"
            
            ref = f"{author_str}. {title}. {journal}. {year}."
            if doi:
                ref += f" doi:{doi}"
            
            references.append(ref)
        
        return "\n".join(references)

    def _summarize_quality_assessments(self, quality_assessments: List[Dict[str, Any]]) -> str:
        """
        Summarize quality assessment results
        """
        if not quality_assessments:
            return "Quality assessment data not available."
        
        total = len(quality_assessments)
        low_risk = sum(1 for qa in quality_assessments if qa.get('risk_of_bias') == 'Low')
        high_risk = sum(1 for qa in quality_assessments if qa.get('risk_of_bias') == 'High')
        
        return f"Of {total} studies assessed, {low_risk} were at low risk of bias, {high_risk} at high risk."

    def export_article_to_word(self, article: Dict[str, Any], filename: str):
        """
        Export article to Word document format
        """
        try:
            from docx import Document
            
            doc = Document()
            
            # Title
            title = article['sections'].get('title', 'Systematic Review Article')
            doc.add_heading(title, 0)
            
            # Abstract
            doc.add_heading('Abstract', 1)
            doc.add_paragraph(article['sections'].get('abstract', ''))
            
            # Main sections
            sections = ['introduction', 'methods', 'results', 'discussion', 'conclusion']
            
            for section in sections:
                if section in article['sections']:
                    doc.add_heading(section.title(), 1)
                    doc.add_paragraph(article['sections'][section])
            
            # References
            if 'references' in article['sections']:
                doc.add_heading('References', 1)
                doc.add_paragraph(article['sections']['references'])
            
            doc.save(filename)
            return True
            
        except ImportError:
            print("python-docx not available for Word export")
            return False
        except Exception as e:
            print(f"Word export failed: {e}")
            return False

    def export_article_to_markdown(self, article: Dict[str, Any], filename: str):
        """
        Export article to Markdown format
        """
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                # Title
                f.write(f"# {article['sections'].get('title', 'Systematic Review Article')}\n\n")
                
                # Abstract
                f.write("## Abstract\n\n")
                f.write(f"{article['sections'].get('abstract', '')}\n\n")
                
                # Main sections
                sections = ['introduction', 'methods', 'results', 'discussion', 'conclusion']
                
                for section in sections:
                    if section in article['sections']:
                        f.write(f"## {section.title()}\n\n")
                        f.write(f"{article['sections'][section]}\n\n")
                
                # References
                if 'references' in article['sections']:
                    f.write("## References\n\n")
                    f.write(f"{article['sections']['references']}\n\n")
            
            return True
            
        except Exception as e:
            print(f"Markdown export failed: {e}")
            return False
